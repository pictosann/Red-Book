import streamlit as st
import os
from openai import OpenAI
import requests
import io
from docx import Document
from docx.shared import Inches

# 页面全局设置
st.set_page_config(page_title="AIGC 自动化流水线 v3.0", layout="wide", page_icon="🚀")

# ================= 侧边栏：系统动态配置 =================
with st.sidebar:
    st.header("⚙️ 系统控制台")
    st.markdown("配置你的 API 密钥（建议在代码中填入默认值）：")

    # 【修复/记忆】Key 隔离存储与默认值设置
    deepseek_key = st.text_input("🔑 DeepSeek API Key", type="password", value=os.getenv("DEEPSEEK_API_KEY", ""))
    silicon_key = st.text_input("🔑 硅基流动 API Key", type="password", value=os.getenv("SILICONFLOW_API_KEY", ""))
    gitee_key = st.text_input("🔑 Gitee AI Token", type="password", value=os.getenv("GITEE_AI_TOKEN", ""))  # 【新增】Gitee Token

    st.divider()

    st.subheader("🎨 模型热切换引擎")
    # 【新增】加入 Z-Image-Turbo
    image_model = st.selectbox(
        "选择文生图底层模型：",
        ("Qwen/Qwen-Image", "Kwai-Kolors/Kolors", "Z-Image-Turbo")
    )
    st.info(f"💡 亮点：本系统已成功实现云端异构算力（SiliconFlow & Gitee）的模型层解耦与热切换！")


# ================= 核心工作流引擎节点 =================

def ask_llm(system_prompt, user_input, api_key):
    """节点 1&2：文本逻辑处理（DeepSeek）"""
    if not api_key: raise Exception("未配置 DeepSeek Key")
    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_input}
    ]
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=messages,
        temperature=0.7
    )
    return response.choices[0].message.content


def generate_image_siliconflow(prompt, model_name, api_key):
    """节点 4a：调用硅基流动 API（返回 URL）"""
    if not api_key: raise Exception("未配置硅基流动 Key")
    url = "https://api.siliconflow.cn/v1/images/generations"
    payload = {
        "model": model_name,
        "prompt": prompt,
        "image_size": "1024x1024",
        "batch_size": 1
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    response = requests.post(url, json=payload, headers=headers)
    result = response.json()
    if "images" in result and len(result["images"]) > 0:
        return result["images"][0]["url"]
    else:
        raise Exception(f"SiliconFlow 画图失败，API 返回: {result}")


def generate_image_gitee(prompt, api_key):
    """节点 4b：【新增】调用 Gitee AI 的 Z-Image-Turbo（返回二进制数据）"""
    if not api_key: raise Exception("未配置 Gitee Token")
    # Gitee 标准推理接口 URL
    url = "https://ai.gitee.com/hf-models/Z-Image-Turbo/api"
    # HuggingFace 标准推理输入格式
    payload = {"inputs": prompt}
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    response = requests.post(url, json=payload, headers=headers)
    # Gitee 推理 API 通常直接返回图片 Bytes 数据
    if response.status_code == 200:
        return response.content
    else:
        raise Exception(f"Gitee 画图失败，代码: {response.status_code}, 返回: {response.text}")


def create_word_doc(final_text, img_data, topic):
    """节点 5：文档打包（在内存中处理图片 BytesIO）"""
    doc = Document()
    doc.add_heading(f'AIGC 自动化内容生成：{topic}', 0)

    # 将图片 Bytes 数据包裹成 Word 能识别的流
    img_stream = io.BytesIO(img_data)
    try:
        doc.add_picture(img_stream, width=Inches(5.0))
    except Exception as e:
        doc.add_paragraph(f"[⚠️ 图片处理失败: {e}]")

    doc.add_paragraph(final_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ================= 主控制台 =================
st.title("🔄 异构多模型内容自动化生成系统")
st.write("演示：【文本生成】 -> 【智能审核】 -> 【画面解析】 -> 【异构画图切换】 -> 【文档打包】")

topic = st.text_input("📝 请输入创作主题：", value="春季男生穿搭")

if st.button("🚀 启动自动化工作流", type="primary"):
    if topic:
        with st.spinner(f"工作流执行中，正在调用 {image_model} ..."):
            try:
                # 1. 初始化需要传递给下游的变量
                img_data_bytes = None

                # --- 节点 1&2：文本与审核（DeepSeek） ---
                # 为了演示快速，我们这里合并调用
                sys_total = "你是一个专业的小红书运营及主编。请根据主题生成一篇高质量、合规、带Emoji和话题标签的图文文案。自我检查如果不符合要求请自行重写。只输出定稿文案。"
                st.toast("节点 1&2 执行中：生成全自动审核定稿...", icon="📝")
                final_text = ask_llm(sys_total, topic, deepseek_key)

                # --- 节点 3：提取画面 ---
                st.toast("节点 3 执行中：提取视觉提示词...", icon="🧠")
                image_prompt = ask_llm(
                    "你是一个AI绘画专家。提取用户文案的核心画面，翻译成一句具体的英文绘图提示词，只输出纯英文。",
                    final_text, deepseek_key)

                # --- 节点 4：异构图像渲染（分支逻辑） ---
                st.toast(f"节点 4 执行中：调用 {image_model} 渲染...", icon="🎨")

                if image_model == "Z-Image-Turbo":
                    # 调用 Gitee，直接获取二进制 Bytes
                    img_data_bytes = generate_image_gitee(image_prompt, gitee_key)
                else:
                    # 调用 SiliconFlow，获取 URL，然后立即下载变成 Bytes
                    img_url = generate_image_siliconflow(image_prompt, image_model, silicon_key)
                    # 立即下载图片数据到内存，实现格式统一
                    st.toast(f"正在从云端下载渲染好的图片...", icon="⬇️")
                    img_data_bytes = requests.get(img_url).content

                st.success("🎉 全链路生产完毕！")

                # --- 展示与文档导出 ---
                st.divider()
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("#### 🖼️ 最终视觉物料")
                    # Streamlit st.image 可以直接显示 Bytes 数据
                    st.image(img_data_bytes, caption=f"模型: {image_model} | Prompt: {image_prompt}",
                             use_container_width=True)
                with col2:
                    st.markdown("#### 📝 定稿文案")
                    st.info(final_text)

                st.divider()
                st.markdown("#### 📥 导出产物")

                # 生成 Word 文档（直接传递我们统一好的 Bytes 数据）
                word_buffer = create_word_doc(final_text, img_data_bytes, topic)
                st.download_button(
                    label="📄 一键下载排版好的 Word 文档",
                    data=word_buffer,
                    file_name=f"AIGC生成_{topic}_{image_model.replace('/', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

            except Exception as e:
                st.error(f"❌ 运行中断，详细错误信息: {e}")
                st.warning("请检查对应的 API Key 是否配置正确，或者接口是否拥堵。")
    else:
        st.error("请输入主题！")
